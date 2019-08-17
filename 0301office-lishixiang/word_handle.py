#!/usr/bin/env Python
# -*- coding:utf-8 -*-
# word_handle.py
# author: lishixiang


from docx import Document
from docx.shared import Pt


"""
1.做一个word操作的类, 可以根据不同的输入内容, 生成word文件, 要求至少完成三类内容的输入: 标题, 副标题, 正文.
"""


class WordHandle():
    "可以操作word的类"
    def __init__(self):
        "创建Document实例"
        self.doc = Document()
        self.content_dic = self.input_content()

    def input_content(self):
        "根据提示让用户输入需要添加到word中的内容"
        content_dic = {}
        content_dic['title'] = input('请输入标题: ')
        content_dic['subtitle'] = input('请输入副标题: ')
        content_dic['paragraph'] = input('请输入正文内容: ')
        bool_ = input('请选择字体是否加粗:\n0 - 不加粗\n其他 - 加粗\n ')
        if bool_ == '0':
            bool_ = False
        else:
            bool_ = True
        content_dic['bold'] = bool_
        content_dic['pt'] = input('请输入字体大小: ')
        content_dic['name'] = input('请输入要保存的文件的名字: ') + '.docx'
        return content_dic

    def word_heading(self):
        "添加word文档的标题"
        self.doc.add_heading(self.content_dic['title'], level=0)

    def word_subtitle(self):
        "添加word文档的副标题"
        self.doc.add_paragraph(self.content_dic['subtitle'], style='Subtitle')

    def word_paragraph(self):
        "添加word文档的正文内容"
        para = self.doc.add_paragraph(style='Body Text')
        para = para.add_run(self.content_dic['paragraph'])
        para.font.bold = self.content_dic['bold']
        para.font.size = Pt(int(self.content_dic['pt']))

    def save_word(self):
        "保存word文档"
        status = {}
        try:
            self.doc.save(self.content_dic['name'])
            status['0'] = 'done'
        except Exception as e:
            status['1'] = e
        return status


def main():
    doc = WordHandle()
    doc.word_heading()
    doc.word_subtitle()
    doc.word_paragraph()
    print(doc.save_word())

if __name__ == "__main__":
    main()